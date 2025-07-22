import os
import hashlib
import datetime
import json
import numpy as np
import faiss
import yaml
from pathlib import Path
from typing import List, Dict, Any
from sentence_transformers import SentenceTransformer


class RecursiveCharacterTextSplitter:
    """优化的文本分割器，避免内存问题"""

    def __init__(self, separators: List[str], chunk_size: int, chunk_overlap: int):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators

    def split_text(self, text: str) -> List[str]:
        """安全分割文本"""
        # 如果文本长度小于块大小，直接返回整个文本
        if len(text) <= self.chunk_size:
            return [text]

        chunks = []
        start = 0
        step = self.chunk_size - self.chunk_overlap

        # 确保步长至少为1
        step = max(1, step)

        while start < len(text):
            end = min(start + self.chunk_size, len(text))
            chunks.append(text[start:end])

            # 移动到下一个块，考虑重叠
            start += step

            # 如果重叠大于0，确保不会回退
            if self.chunk_overlap > 0 and start < end:
                start = end - self.chunk_overlap

        return chunks


class VectorStoreBuilder:
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.embedding_model = self._init_embedding_model()
        self.text_splitter = self._init_text_splitter()

    def _init_embedding_model(self) -> SentenceTransformer:
        """初始化本地 BGE-m3 嵌入模型"""
        # 处理可能的嵌套 embedding 配置
        embedding_config = self.config.get('embedding')
        if not embedding_config:
            # 如果顶层没有 embedding，检查 processing 下是否有
            if 'processing' in self.config and 'embedding' in self.config['processing']:
                embedding_config = self.config['processing']['embedding']
            else:
                raise ValueError("配置文件中缺少 'embedding' 部分")

        model_path = embedding_config.get('model_path', embedding_config.get('model_name'))
        if not model_path:
            raise ValueError("embedding 配置中缺少 model_path 或 model_name")

        device = embedding_config.get('device', 'cpu')

        print(f"初始化嵌入模型: {model_path} (设备: {device})")
        return SentenceTransformer(
            model_name_or_path = model_path,
            device = device
        )

    def _init_text_splitter(self):
        """初始化文本分割器"""
        processing = self.config['processing']
        print(f"初始化文本分割器: 块大小={processing['chunk_size']}, 重叠={processing['chunk_overlap']}")
        return RecursiveCharacterTextSplitter(
            separators = processing['separators'],
            chunk_size = processing['chunk_size'],
            chunk_overlap = processing['chunk_overlap']
        )

    def _extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """提取文件元数据"""
        stat = file_path.stat()
        metadata_config = self.config.get('metadata', {})

        metadata = {
            "source":file_path.name,  # 默认使用文件名
            "file_id":hashlib.md5(str(file_path).encode()).hexdigest()[:8],
            "file_size":stat.st_size,
            "created_time":datetime.datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified_time":datetime.datetime.fromtimestamp(stat.st_mtime).isoformat(),
            "file_type":file_path.suffix.lower()[1:],
            "title":metadata_config.get('default_title', '')
        }

        # 如果配置了source_field，则覆盖source
        if 'source_field' in metadata_config:
            if metadata_config['source_field'] == 'file_name':
                metadata['source'] = file_path.name
            # 其他字段可以扩展

        return metadata

    def _read_file(self, file_path: Path) -> str:
        """读取文件内容"""
        # 确保路径是绝对路径
        abs_path = file_path.resolve()
        print(f"读取文件: {abs_path}")

        # PDF文件
        if file_path.suffix.lower() == '.pdf':
            try:
                import PyPDF2
                print(f"尝试读取PDF文件: {abs_path}")
                with open(abs_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    for i, page in enumerate(reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text += f"第 {i + 1} 页:\n{page_text}\n\n"
                        except Exception as e:
                            print(f"提取PDF第 {i + 1} 页内容时出错: {str(e)}")
                    print(f"PDF文件内容长度: {len(text)}")
                    return text
            except ImportError:
                print("错误: PyPDF2未安装，无法处理PDF文件。请运行: pip install PyPDF2")
                return ""
            except Exception as e:
                # 获取详细的错误信息
                import traceback
                error_msg = traceback.format_exc()
                print(f"处理PDF文件 {abs_path} 时出错:\n{error_msg}")
                return ""

        # Word文档
        if file_path.suffix.lower() in ['.docx']:
            try:
                from docx import Document
                print(f"尝试读取Word文档: {abs_path}")
                doc = Document(abs_path)

                # 提取所有段落文本
                content = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        content.append(para.text)

                # 提取表格文本
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if para.text.strip():
                                    content.append(para.text)

                full_content = '\n'.join(content)
                print(f"成功读取Word文档，内容长度: {len(full_content)}")
                return full_content
            except ImportError:
                print("错误: python-docx未安装，无法处理Word文档。请运行: pip install python-docx")
                return ""
            except Exception as e:
                # 获取详细的错误信息
                import traceback
                error_msg = traceback.format_exc()
                print(f"处理Word文档 {abs_path} 时出错:\n{error_msg}")
                return ""

        # 其他文件类型跳过
        print(f"不支持的文件类型: {file_path.suffix}")
        return ""

    def _safe_split_text(self, text: str) -> List[str]:
        """安全的分割方法，避免内存问题"""
        chunk_size = self.text_splitter.chunk_size
        chunk_overlap = self.text_splitter.chunk_overlap

        # 如果文本长度小于块大小，直接返回整个文本
        if len(text) <= chunk_size:
            return [text]

        chunks = []
        start = 0
        step = max(1, chunk_size - chunk_overlap)  # 确保步长至少为1

        while start < len(text):
            end = min(start + chunk_size, len(text))
            chunks.append(text[start:end])
            start += step

        return chunks

    def process_file(self, file_path: Path) -> List[Dict[str, Any]]:
        """处理单个文件"""
        print(f"处理文件: {file_path}")
        try:
            content = self._read_file(file_path)
            if not content:
                print(f"文件内容为空: {file_path}")
                return []
            elif not content.strip():
                print(f"文件内容全为空白: {file_path}")
                return []

            # 检查文件大小，如果太大则警告
            content_size = len(content.encode('utf-8'))
            print(f"文件内容大小: {content_size} 字节")

            base_meta = self._extract_metadata(file_path)

            try:
                chunks = self.text_splitter.split_text(content)
            except MemoryError:
                print(f"内存不足，尝试简化分割: {file_path}")
                chunks = self._safe_split_text(content)
            except Exception as e:
                print(f"分割文本时出错: {str(e)}")
                chunks = self._safe_split_text(content)

            print(f"生成 {len(chunks)} 个文本块")

            results = []
            for i, chunk in enumerate(chunks):
                # 确保块大小合理
                if len(chunk) > self.text_splitter.chunk_size * 2:
                    print(f"警告: 块 {i} 过大 ({len(chunk)} 字符)")

                results.append({
                    "id":f"{base_meta['file_id']}-{i}",
                    "content":chunk,
                    "metadata":{**base_meta, "chunk_index":i}
                })

            return results
        except Exception as e:
            # 获取详细的错误信息
            import traceback
            error_msg = traceback.format_exc()
            print(f"处理文件 {file_path} 时严重出错:\n{error_msg}")
            return []

    def build_vector_store(self):
        """构建向量存储"""
        data_config = self.config['data']

        # 获取所有文件
        source_dir = Path(data_config['input_path'])
        file_patterns = data_config['file_pattern']

        print(f"扫描目录: {source_dir.resolve()}")
        print(f"文件模式: {file_patterns}")

        file_paths = []
        for pattern in file_patterns:
            found = list(source_dir.glob(pattern))
            print(f"模式 '{pattern}' 找到 {len(found)} 个文件")
            file_paths.extend(found)

        if not file_paths:
            print(f"在 {source_dir} 中未找到匹配文件")
            return None

        print(f"找到 {len(file_paths)} 个匹配文件")

        # 处理所有文件
        all_chunks = []
        for file_path in file_paths:
            if file_path.is_file():
                print(f"\n开始处理文件: {file_path}")
                chunks = self.process_file(file_path)
                if chunks:
                    print(f"成功处理: {file_path.name} -> 生成 {len(chunks)} 个文本块")
                    all_chunks.extend(chunks)
                else:
                    print(f"文件未生成文本块: {file_path.name}")

        if not all_chunks:
            print("未提取到有效文本块")
            return None

        print(f"\n总共处理了 {len(all_chunks)} 个文本块")

        # 提取文本和元数据
        texts = [chunk["content"] for chunk in all_chunks]
        metadata = [chunk["metadata"] for chunk in all_chunks]

        # 生成嵌入向量
        print(f"\n为 {len(texts)} 个文本块生成嵌入...")

        # 处理可能的嵌套 embedding 配置
        embedding_config = self.config.get('embedding')
        if not embedding_config:
            if 'processing' in self.config and 'embedding' in self.config['processing']:
                embedding_config = self.config['processing']['embedding']
            else:
                raise ValueError("配置文件中缺少 'embedding' 部分")

        encode_kwargs = embedding_config.get('encode_kwargs', {})

        # 分批处理避免内存问题
        batch_size = embedding_config.get('batch_size', 32)
        embeddings = np.zeros((len(texts), 1024), dtype = np.float32)  # 假设维度为1024

        for i in range(0, len(texts), batch_size):
            batch_texts = texts[i:i + batch_size]
            print(f"处理批次: {i // batch_size + 1}/{len(texts) // batch_size + 1}")

            batch_embeddings = self.embedding_model.encode(
                batch_texts,
                batch_size = batch_size,
                show_progress_bar = False,
                normalize_embeddings = encode_kwargs.get('normalize_embeddings', True)
            )

            # 确保嵌入维度一致
            if batch_embeddings.shape[1] != embeddings.shape[1]:
                # 如果维度不匹配，调整嵌入矩阵
                embeddings = np.zeros((len(texts), batch_embeddings.shape[1]), dtype = np.float32)

            embeddings[i:i + batch_size] = batch_embeddings

        # 创建FAISS索引
        dimension = embeddings.shape[1]
        index = faiss.IndexFlatL2(dimension)
        index.add(embeddings)

        # 保存向量索引和元数据
        output_db = Path(data_config['output_db'])
        output_db.parent.mkdir(parents = True, exist_ok = True)

        # 保存FAISS索引
        faiss.write_index(index, str(output_db))

        # 保存元数据
        metadata_file = output_db.with_suffix('.json')
        with open(metadata_file, 'w', encoding = 'utf-8') as f:
            json.dump(metadata, f, ensure_ascii = False, indent = 2)

        print(f"\n向量存储构建完成!")
        print(f"FAISS索引文件: {output_db}")
        print(f"元数据文件: {metadata_file}")
        print(f"总文档块数: {len(metadata)}")
        print(f"嵌入维度: {dimension}")

        return {
            "index_file":str(output_db),
            "metadata_file":str(metadata_file),
            "total_chunks":len(metadata),
            "embedding_dim":dimension
        }


def load_config(config_path: str) -> Dict[str, Any]:
    """从YAML文件加载配置"""
    if not Path(config_path).exists():
        print(f"配置文件 {config_path} 不存在!")
        return None

    print(f"加载配置文件: {Path(config_path).resolve()}")
    with open(config_path, 'r', encoding = 'utf-8') as f:
        config = yaml.safe_load(f)

    # 修复嵌套的 embedding 配置
    if 'processing' in config and 'embedding' in config['processing']:
        print("检测到嵌套的 embedding 配置，正在修复...")
        config['embedding'] = config['processing']['embedding']
        del config['processing']['embedding']

    return config


if __name__ == "__main__":
    # 加载配置文件
    config_path = "faiss_config.yaml"
    config = load_config(config_path)

    if not config:
        print("无法加载配置，程序退出")
        exit(1)

    print("\n配置加载成功:")
    print(f"输入路径: {config['data']['input_path']}")
    print(f"文件模式: {config['data']['file_pattern']}")
    print(f"输出路径: {config['data']['output_db']}")

    # 检查 embedding 配置是否存在
    if 'embedding' in config:
        model_path = config['embedding'].get('model_path', config['embedding'].get('model_name'))
        print(f"模型路径: {model_path}")
        print(f"设备: {config['embedding'].get('device', 'cpu')}")
    else:
        print("警告: 配置文件中缺少 'embedding' 部分!")

    print("\n开始构建向量存储...")
    builder = VectorStoreBuilder(config)
    result = builder.build_vector_store()

    if result:
        print("\n向量库构建成功!")
        print(f"FAISS索引已保存到: {result['index_file']}")
        print(f"元数据已保存到: {result['metadata_file']}")
    else:
        print("\n向量库构建失败!")
        print("可能原因:")
        print("1. 文档目录中没有匹配的文件")
        print("2. 文档内容为空或格式不支持")
        print("3. 文件读取过程中出错")